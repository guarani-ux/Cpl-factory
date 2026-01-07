import streamlit as st
import pandas as pd
from docx import Document
import io
import PyPDF2

st.set_page_config(page_title="CPL Factory Pro", layout="wide")

# --- CORE LOGIC: DOCUMENT PARSING ---
def extract_text(uploaded_file):
    if uploaded_file.type == "application/pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        return " ".join([page.extract_text() for page in reader.pages])
    return str(uploaded_file.read(), "utf-8")

st.title("🚀 CPL Factory: Production Architect")

with st.sidebar:
    st.header("1. Upload Brief")
    doc = st.file_uploader("Upload CPL Strategy Doc", type=["pdf", "txt"])
    st.header("2. Algorithm Settings")
    tone = st.selectbox("Tone", ["Executive", "Casual/Viral", "Technical", "Educational"])
    platform = st.radio("Primary Distribution", ["YouTube Main", "Shorts/TikTok", "LinkedIn Video"])

if doc:
    raw_text = extract_text(doc)
    st.success("Brief Analyzed. Ready to Architect.")

    if st.button("Generate Full Production Suite"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📺 Video Script (3-Act Structure)")
            script = f"""
            **THE HOOK (0:00-0:15):** "Most people think {tone} strategy is about volume, but for CPL, it's about context. Here is why..."
            
            **THE VALUE (0:15-4:00):** - Point 1: Extracted from your brief...
            - Point 2: Algorithm-optimized delivery...
            
            **THE CTA (4:00-End):** "Download the full CPL guide in the description below."
            """
            st.markdown(script)
            
            # Export to Word
            doc_out = Document()
            doc_out.add_heading('CPL Video Script', 0)
            doc_out.add_paragraph(script)
            bio = io.BytesIO()
            doc_out.save(bio)
            st.download_button("📥 Download Script (.docx)", bio.getvalue(), "CPL_Script.docx")

        with col2:
            st.subheader("📋 Asana Production Map")
            tasks = [
                {"Task": "Finalize Retention-Focused Script", "Section": "Pre-Prod", "Due": "Day 1"},
                {"Task": "A/B Test Thumbnail Concept 1 vs 2", "Section": "Creative", "Due": "Day 2"},
                {"Task": "Film A-Roll (Executive Setup)", "Section": "Production", "Due": "Day 4"},
                {"Task": "B-Roll Overlay & Motion Graphics", "Section": "Editing", "Due": "Day 6"},
                {"Task": "SEO Metadata & Chapter Markers", "Section": "Post-Prod", "Due": "Day 8"},
                {"Task": "Compliance & Brand Safety Check", "Section": "Legal", "Due": "Day 9"}
            ]
            df = pd.DataFrame(tasks)
            st.table(df)
            st.download_button("📥 Download Asana CSV", df.to_csv(index=False), "asana_plan.csv")

else:
    st.info("Upload a document in the sidebar to unlock the Production Architect.")
